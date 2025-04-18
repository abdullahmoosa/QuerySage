import streamlit as st
from datetime import datetime
import uuid
import psycopg2
from pinecone import Pinecone
from langchain.chains import RetrievalQA
import openai
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Pinecone as LangchainPinecone
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from dotenv import load_dotenv
import tempfile
import logging
import json
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document
import io
from urllib.parse import quote_plus
from supabase import create_client, Client

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

load_dotenv()
openai.api_key = os.environ["OPENAI_API_KEY"]
PINECONE_API_KEY = os.environ["PINECONE_API_KEY"]
INDEX_NAME = "courseassesmentsystem"

# Extract Supabase URL and key from the DATABASE_URL
# Format: postgresql://postgres:password@aws-0-us-west-1.pooler.supabase.com:6543/postgres?options=reference%3Dtmwetgegqdxnibsmwrxy
db_url = os.environ.get("DATABASE_URL", "")
# The project reference is in the options parameter
supabase_ref = db_url.split("reference%3D")[-1] if "reference%3D" in db_url else ""
# Construct Supabase URL and key (these should ideally be in .env)
SUPABASE_URL = f"https://{supabase_ref}.supabase.co"
# For security, the anon key should be in .env, but we'll use a placeholder for now
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InRtd2V0Z2VncWR4bmliczt3cnh5Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3MTU5NTI2MDAsImV4cCI6MjAzMTUyODYwMH0.placeholder")
# Get the service role key if available (has admin privileges)
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")

# Global variables for database and Pinecone connections
conn = None
cursor = None
index = None
supabase = None
supabase_admin = None

# Initialize Pinecone
def init_pinecone():
    pc = Pinecone(api_key=PINECONE_API_KEY)
    logger.info("Initializing Pinecone connection")

    # Create index if it doesn't exist
    if INDEX_NAME not in pc.list_indexes().names():
        pc.create_index(
            name=INDEX_NAME,
            dimension=1536,
            metric='cosine'
        )
        logger.info(f"Successfully created Pinecone index: {INDEX_NAME}")
    return pc.Index(INDEX_NAME)

# Initialize Supabase
def init_supabase():
    try:
        logger.info(f"Initializing Supabase connection to {SUPABASE_URL}")
        
        # Create client with anon key for regular operations
        sb = create_client(SUPABASE_URL, SUPABASE_KEY)
        
        # If service role key is available, create an admin client
        sb_admin = None
        if SUPABASE_SERVICE_KEY:
            logger.info("Service role key found, creating admin client")
            sb_admin = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
        
        logger.info("Successfully connected to Supabase")
        return sb, sb_admin
    except Exception as e:
        logger.error(f"Failed to connect to Supabase: {str(e)}")
        raise

# Initialize PostgreSQL
def init_postgres():
    try:
        # Construct the connection string with the encoded password
        connection_string = os.environ.get("DATABASE_URL")        
        new_conn = psycopg2.connect(connection_string)
        new_cursor = new_conn.cursor()
        logger.info("Successfully connected to PostgreSQL database")
        return new_conn, new_cursor
    except Exception as e:
        logger.error(f"Failed to connect to PostgreSQL: {str(e)}")
        raise

user_id = "0194b305-f06e-7661-810e-8f35c12ab058"

def process_pdf(file_content) -> str:
    """Process PDF file and return its text content"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        text_content = []
        
        for page in pdf_reader.pages:
            text_content.append(page.extract_text())
            
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        raise

def docx_to_markdown(docx_path):
    """Convert a DOCX file to markdown format"""
    try:
        doc = Document(docx_path)
        markdown_content = []

        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Check for heading style
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                    markdown_content.append('#' * level + ' ' + paragraph.text)
                else:
                    # Handle text formatting
                    text = paragraph.text
                    
                    # Check for bold, italic, etc.
                    for run in paragraph.runs:
                        if run.bold:
                            text = text.replace(run.text, f"**{run.text}**")
                        if run.italic:
                            text = text.replace(run.text, f"*{run.text}*")
                        if run.underline:
                            text = text.replace(run.text, f"__{run.text}__")
                    
                    markdown_content.append(text)

        # Process tables
        for table in doc.tables:
            table_rows = []
            
            # Process header row
            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            table_rows.append("| " + " | ".join(header_row) + " |")
            
            # Add separator row
            separator = "|" + "|".join(["---"] * len(header_row)) + "|"
            table_rows.append(separator)
            
            # Process data rows
            for row in table.rows[1:]:
                row_cells = []
                for cell in row.cells:
                    row_cells.append(cell.text.strip())
                table_rows.append("| " + " | ".join(row_cells) + " |")
            
            # Add table to markdown content
            markdown_content.append("\n".join(table_rows))

        # Handle lists (basic implementation)
        # This is a simplified approach; a more robust implementation would track list levels
        
        return "\n\n".join(markdown_content)
    except Exception as e:
        logger.error(f"Error converting DOCX to markdown: {str(e)}")
        # Return a simple text extraction as fallback
        try:
            doc = Document(docx_path)
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except:
            return "Error processing template document"

def process_docx(file_content) -> str:
    """Process DOCX file and return its text content"""
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
                
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        raise

def process_txt(file_content) -> str:
    """Process TXT file and return its text content"""
    try:
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        # Try different encodings if UTF-8 fails
        try:
            return file_content.decode('latin-1')
        except Exception as e:
            logger.error(f"Error processing TXT: {str(e)}")
            raise

def upload_template_to_supabase(file, question_id: str) -> str:
    """
    Upload a template file to Supabase storage
    
    Args:
        file: The uploaded file
        question_id: UUID of the question
    
    Returns:
        str: The URL of the uploaded file
    """
    try:
        logger.info(f"Uploading template file {file.name} for question {question_id}")
        
        # Create a unique file name to avoid collisions
        file_extension = os.path.splitext(file.name)[1].lower()
        unique_filename = f"template_{question_id}{file_extension}"
        
        # Upload to Supabase storage
        file_content = file.getvalue()
        bucket_name = "templates"
        
        # Ensure the bucket exists with proper permissions
        ensure_storage_bucket_exists(bucket_name, public=True)
        
        # Try to upload the file
        try:
            # Upload the file
            supabase.storage.from_(bucket_name).upload(
                path=unique_filename,
                file=file_content,
                file_options={"content-type": file.type}
            )
            
            # Get the public URL
            file_url = supabase.storage.from_(bucket_name).get_public_url(unique_filename)
            logger.info(f"Successfully uploaded template file to {file_url}")
            
            return file_url
        except Exception as upload_error:
            logger.error(f"Error uploading file: {str(upload_error)}")
            raise Exception(f"Failed to upload file: {str(upload_error)}")
    except Exception as e:
        logger.error(f"Failed to upload template file: {str(e)}")
        raise

def process_uploaded_file(uploaded_file):
    """Process an uploaded file and return its text content"""
    logger.info(f"Processing uploaded file: {uploaded_file.name}")
    
    try:
        file_content = uploaded_file.getvalue()
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension == '.pdf':
            text_content = process_pdf(file_content)
        elif file_extension == '.docx':
            text_content = process_docx(file_content)
        elif file_extension == '.txt':
            text_content = process_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        logger.info(f"Successfully processed file: {uploaded_file.name}")
        return text_content
    except Exception as e:
        logger.error(f"Error processing file {uploaded_file.name}: {str(e)}")
        raise


def split_text(text: str, chunk_size: int = 2000, chunk_overlap: int = 200) -> List[str]:
    """Split text into smaller chunks"""
    logger.debug(f"Splitting text with chunk_size={chunk_size}, overlap={chunk_overlap}")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    chunks = text_splitter.split_text(text)
    logger.debug(f"Split text into {len(chunks)} chunks")
    return chunks


def create_vector_store(documents: List[str], namespace: str) -> LangchainPinecone:
    """Create a Pinecone vector store with the specified namespace"""
    logger.info(f"Creating vector store in namespace: {namespace}")
    
    try:
        embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
        
        # Create vector store
        vector_store = LangchainPinecone.from_existing_index(
            index_name=INDEX_NAME,
            embedding=embeddings,
            namespace=namespace
        )
        logger.info(f"Successfully created vector store in namespace: {namespace}")
        
        total_chunks = 0
        # Process each document
        for doc_index, doc in enumerate(documents, 1):
            # Split the document into smaller chunks
            chunks = split_text(doc)
            total_chunks += len(chunks)
            
            logger.info(f"Processing document {doc_index}/{len(documents)} with {len(chunks)} chunks")
            
            # Add chunks in batches
            batch_size = 50
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                try:
                    vector_store.add_texts(batch)
                    logger.info(f"Successfully processed batch {i//batch_size + 1} in document {doc_index}")
                    # Add a small progress indicator
                    st.write(f"Processed {min(i + batch_size, len(chunks))}/{len(chunks)} chunks...")
                except Exception as e:
                    logger.error(f"Error processing batch {i//batch_size + 1} in document {doc_index}: {str(e)}")
                    continue
        
        logger.info(f"Completed vector store creation. Total chunks processed: {total_chunks}")
        return vector_store
    except Exception as e:
        logger.error(f"Failed to create vector store: {str(e)}")
        raise


def process_file_uploads(files, question_id: str, content_type: str) -> tuple[str | None, list]:
    """
    Process uploaded files and create vector store
    
    Args:
        files: The uploaded files
        question_id: UUID of the question
        content_type: Type of content ('guidelines' or 'sample_answers')
    
    Returns:
        tuple: (vector_db_url, processed_documents)
    """
    if not files:
        return None, []

    documents = []
    for file in files:
        try:
            content = process_uploaded_file(file)
            documents.append(content)
        except Exception as e:
            st.error(f"Error processing {content_type} file {file.name}: {str(e)}")
            continue

    if documents:
        namespace = f"question_{question_id}_{content_type}"
        vector_store = create_vector_store(documents, namespace)
        return namespace, documents
    
    return None, []


def handle_content_input(
    input_method: str,
    files,
    text_content: str,
    question_id: str,
    content_type: str
) -> tuple[str | None, str | None]:
    """
    Handle both file and text input methods
    
    Args:
        input_method: "Upload Files" or "Text Input"
        files: Uploaded files (if any)
        text_content: Text input content (if any)
        question_id: UUID of the question
        content_type: Type of content ('guidelines' or 'sample_answers')
    
    Returns:
        tuple: (vector_db_url, text_content)
    """
    if input_method == "Upload Files":
        logger.info(f"Processing {content_type} as file upload")
        vector_db_url, _ = process_file_uploads(files, question_id, content_type)
        return vector_db_url, None
    else:
        logger.info(f"Processing {content_type} as text input")
        return None, text_content


def execute_db_query(query: str, params: tuple = None, fetch: bool = False) -> Any:
    """Execute a database query with logging"""
    try:
        logger.debug(f"Executing query: {query} with params: {params}")
        cursor.execute(query, params)
        
        if fetch:
            result = cursor.fetchall()
            logger.debug(f"Query returned {len(result)} rows")
            return result
        else:
            conn.commit()
            logger.debug("Query executed successfully")
            return None
    except Exception as e:
        logger.error(f"Database error: {str(e)}")
        conn.rollback()
        raise


def create_answer_doc(answer: str, filename: str):
    """Create a document with just the answer"""
    doc = Document()
    doc.add_heading('Generated Answer', 0)
    doc.add_paragraph(answer)
    doc.save(filename)

def create_qa_doc(question: str, answer: str, filename: str):
    """Create a document with both question and answer"""
    doc = Document()
    doc.add_heading('Question and Answer', 0)
    
    # Add question section
    doc.add_heading('Question:', level=1)
    doc.add_paragraph(question)
    
    # Add answer section
    doc.add_heading('Answer:', level=1)
    doc.add_paragraph(answer)
    
    doc.save(filename)

def create_all_qa_doc(questions_and_answers: list, filename: str):
    """Create a document with multiple questions and answers"""
    doc = Document()
    doc.add_heading('All Questions and Answers', 0)
    
    for idx, (question, answer) in enumerate(questions_and_answers, 1):
        # Add section for each Q&A pair
        doc.add_heading(f'Question {idx}:', level=1)
        doc.add_paragraph(question)
        
        doc.add_heading('Answer:', level=2)
        doc.add_paragraph(answer)
        
        # Add a page break after each Q&A except the last one
        if idx < len(questions_and_answers):
            doc.add_page_break()
    
    doc.save(filename)

def generate_all_answers(subject_id: str, embeddings, llm) -> list:
    """Generate answers for all questions in a subject"""
    # Get all questions for the subject
    questions = execute_db_query(
        """SELECT q.*, s.vector_db_url as subject_vector_db_url 
           FROM questions q 
           JOIN subjects s ON q.subject_id = s.id 
           WHERE q.subject_id = %s 
           ORDER BY q.created_at""",
        (subject_id,),
        fetch=True
    )
    
    qa_pairs = []
    for q in questions:
        logger.info(f"Generating answer for question: {q[1][:100]}...")
        
        # Initialize content list
        retrieved_content = []
        
        # 1. Get subject-level information
        subject_vector_store = LangchainPinecone.from_existing_index(
            index_name=INDEX_NAME,
            embedding=embeddings,
            namespace=q[-1]  # subject_vector_db_url from join
        )
        
        # Get subject content
        subject_retriever = subject_vector_store.as_retriever(search_kwargs={"k": 5})
        subject_docs = subject_retriever.get_relevant_documents(q[1])
        retrieved_content.extend([doc.page_content for doc in subject_docs])
        
        # 2. Get question-specific guidelines
        is_guideline_textbox = q[4]  # is_guideline_textbox
        guideline_text = q[3]  # guideline_text_box
        guideline_vector_db_url = q[2]  # guideline_vector_db_url
        
        if is_guideline_textbox and guideline_text:
            retrieved_content.append(f"Question Guidelines:\n{guideline_text}")
        elif not is_guideline_textbox and guideline_vector_db_url:
            try:
                question_vector_store = LangchainPinecone.from_existing_index(
                    index_name=INDEX_NAME,
                    embedding=embeddings,
                    namespace=guideline_vector_db_url
                )
                guideline_retriever = question_vector_store.as_retriever(search_kwargs={"k": 3})
                guideline_docs = guideline_retriever.get_relevant_documents(q[1])
                retrieved_content.extend([doc.page_content for doc in guideline_docs])
            except Exception as e:
                logger.warning(f"Could not use question guidelines: {str(e)}")
        
        # Get additional context
        additional_context = []
        
        # 3. Handle sample answers based on input type
        is_sample_answer_textbox = q[7]  # is_sample_answer_textbox
        if is_sample_answer_textbox:
            sample_answers_text = q[6]  # sample_answers_textbox
            if sample_answers_text:
                additional_context.append(f"Sample Answers:\n{sample_answers_text}")
        else:
            sample_answers_vector_db_url = q[5]  # sample_answers_vector_db_url
            if sample_answers_vector_db_url:
                try:
                    sample_answers_store = LangchainPinecone.from_existing_index(
                        index_name=INDEX_NAME,
                        embedding=embeddings,
                        namespace=sample_answers_vector_db_url
                    )
                    sample_answers_retriever = sample_answers_store.as_retriever(search_kwargs={"k": 3})
                    sample_answers_docs = sample_answers_retriever.get_relevant_documents(q[1])
                    if sample_answers_docs:
                        additional_context.append("Sample Answers:\n" + "\n".join([doc.page_content for doc in sample_answers_docs]))
                except Exception as e:
                    logger.warning(f"Could not use sample answers from vector store: {str(e)}")
        
        if q[8]:  # Instructions
            additional_context.append(f"Instructions:\n{q[8]}")
        
        # 4. Handle template if available
        template_url = q[9]  # template_url
        template_content = None
        if template_url:
            try:
                logger.info(f"Found template at {template_url}, downloading...")
                # Extract filename from URL
                filename = template_url.split('/')[-1]
                bucket_name = "templates"
                
                # Use admin client if available for download operations
                storage_client = supabase_admin if supabase_admin else supabase
                
                # Download the file
                template_data = storage_client.storage.from_(bucket_name).download(filename)
                
                # Convert to markdown
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_file.write(template_data)
                    tmp_file_path = tmp_file.name
                
                template_content = docx_to_markdown(tmp_file_path)
                
                # Clean up temporary file
                try:
                    os.unlink(tmp_file_path)
                except Exception as e:
                    logger.warning(f"Error cleaning up temporary file: {str(e)}")
                
                logger.info("Successfully processed template")
                additional_context.append(f"Template Structure:\n{template_content}")
            except Exception as e:
                logger.error(f"Failed to process template: {str(e)}")
        
        # Prepare prompt and generate answer
        prompt = """
            Question: {}

            Retrieved Information:
            {}

            {}

            {}

            You are an expert in the subject of the question. Please provide a comprehensive answer to the question using the retrieved information above.""".format(
            q[1],
            "-" * 50 + "\n" + "\n".join(retrieved_content) + "\n" + "-" * 50,
            "\n\n".join(additional_context),
            "Follow the template structure provided above when formatting your answer." if template_content else ""
        )
        
        answer = llm.predict(prompt)
        qa_pairs.append((q[1], answer))
        logger.info("Answer generated successfully")
    
    return qa_pairs

def delete_vector_store_namespace(namespace: str):
    """Delete a namespace from the Pinecone vector store"""
    try:
        logger.info(f"Starting deletion of namespace '{namespace}' from vector store")
        index.delete(namespace=namespace, delete_all=True)
        logger.info(f"✓ Successfully deleted vector store namespace: '{namespace}'")
    except Exception as e:
        logger.error(f"❌ Failed to delete vector store namespace '{namespace}': {str(e)}")
        raise

def delete_question(question_id: str):
    """Delete a question and its associated vector store data"""
    try:
        # Get question data before deletion
        question = execute_db_query(
            """SELECT id, text, guideline_vector_db_url, guideline_text_box, 
                     is_guideline_textbox::boolean, sample_answers_vector_db_url, 
                     sample_answers_textbox, instructions, created_at, subject_id, template_url 
              FROM questions 
              WHERE id = %s""",
            (question_id,),
            fetch=True
        )[0]
        
        logger.info("\n=== Starting deletion of question ===")
        logger.info(f"Question ID: {question_id}")
        logger.info(f"Question Text: {question[1][:100]}...")
        
        # Debug log to show actual is_guideline_textbox value
        logger.info(f"Raw is_guideline_textbox value from DB: {question[4]} (type: {type(question[4])})")
        logger.info(f"Raw is_guideline_textbox value repr: {repr(question[4])}")
        
        # Properly check is_guideline_textbox value
        is_guideline_textbox = bool(question[4])  # Convert to boolean to ensure proper type
        logger.info(f"Converted is_guideline_textbox value: {is_guideline_textbox} (type: {type(is_guideline_textbox)})")
        input_method = "Text Input" if is_guideline_textbox else "File Upload"
        logger.info(f"Input Method: {input_method}")
        
        # Check for any vector stores to delete
        vector_stores_deleted = False
        
        # For file upload questions (is_guideline_textbox = False), check and delete vector stores
        if not is_guideline_textbox:  # Using the properly converted boolean value
            logger.info("Question uses file upload - checking for vector stores")
            
            # Check and delete guideline vector store
            if question[2]:  # guideline_vector_db_url
                logger.info(f"Found guideline vector store to delete: {question[2]}")
                try:
                    delete_vector_store_namespace(question[2])
                    logger.info("✓ Guidelines vector store deleted")
                    vector_stores_deleted = True
                except Exception as e:
                    logger.error(f"Failed to delete guideline vector store: {str(e)}")
            else:
                logger.info("No guideline vector store found")
            
            # Check and delete sample answers vector store
            if question[5]:  # sample_answers_vector_db_url
                logger.info(f"Found sample answers vector store to delete: {question[5]}")
                try:
                    delete_vector_store_namespace(question[5])
                    logger.info("✓ Sample answers vector store deleted")
                    vector_stores_deleted = True
                except Exception as e:
                    logger.error(f"Failed to delete sample answers vector store: {str(e)}")
            else:
                logger.info("No sample answers vector store found")
        else:
            logger.info("Question uses text input - no vector stores to delete")
        
        if not vector_stores_deleted:
            logger.info("No vector stores were deleted")
        
        # Check and delete template file from Supabase storage
        template_url = question[10]  # template_url
        if template_url:
            logger.info(f"Found template file to delete: {template_url}")
            try:
                # Extract filename from URL
                filename = template_url.split('/')[-1]
                bucket_name = "templates"
                
                # Use admin client if available for delete operations
                storage_client = supabase_admin if supabase_admin else supabase
                
                # Delete the file
                storage_client.storage.from_(bucket_name).remove([filename])
                logger.info("✓ Template file deleted from Supabase storage")
            except Exception as e:
                logger.error(f"Failed to delete template file: {str(e)}")
        
        # Delete from database
        execute_db_query(
            "DELETE FROM questions WHERE id = %s",
            (question_id,)
        )
        logger.info("✓ Question record deleted from database")
        logger.info("=== Question deletion completed ===\n")
        return True
    except Exception as e:
        logger.error(f"❌ Error deleting question {question_id}: {str(e)}")
        return False

def delete_subject(subject_id: str):
    """Delete a subject and all its questions"""
    try:
        # Get subject data before deletion
        subject = execute_db_query(
            "SELECT * FROM subjects WHERE id = %s",
            (subject_id,),
            fetch=True
        )[0]
        
        logger.info(f"\n=== Starting deletion of subject ===")
        logger.info(f"Subject ID: {subject_id}")
        logger.info(f"Subject Name: {subject[1]}")
        
        # Get all questions for this subject
        questions = execute_db_query(
            "SELECT id FROM questions WHERE subject_id = %s",
            (subject_id,),
            fetch=True
        )
        
        # Delete all questions first
        logger.info(f"Found {len(questions)} questions to delete")
        for i, question in enumerate(questions, 1):
            logger.info(f"\nDeleting question {i} of {len(questions)}")
            if delete_question(question[0]):
                logger.info(f"✓ Successfully deleted question {i}")
            else:
                logger.warning(f"⚠️ Failed to delete question {i}")
        
        # Delete subject's vector store data if exists
        if subject[3]:  # vector_db_url
            logger.info(f"Found subject vector store to delete: {subject[3]}")
            delete_vector_store_namespace(subject[3])
            logger.info("✓ Subject vector store deleted")
        else:
            logger.info("No subject vector store to delete")
        
        # Delete from database
        execute_db_query(
            "DELETE FROM subjects WHERE id = %s",
            (subject_id,)
        )
        logger.info("✓ Subject record deleted from database")
        logger.info("=== Subject deletion completed ===\n")
        return True
    except Exception as e:
        logger.error(f"❌ Error deleting subject {subject_id}: {str(e)}")
        return False

def delete_course(course_id: str):
    """Delete a course and all its subjects"""
    try:
        # Get course data first
        course = execute_db_query(
            "SELECT * FROM courses WHERE id = %s",
            (course_id,),
            fetch=True
        )[0]
        
        logger.info(f"\n=== Starting deletion of course ===")
        logger.info(f"Course ID: {course_id}")
        logger.info(f"Course Name: {course[1]}")
        
        # Get all subjects for this course
        subjects = execute_db_query(
            "SELECT id, name FROM subjects WHERE course_id = %s",
            (course_id,),
            fetch=True
        )
        
        # Delete all subjects first
        logger.info(f"Found {len(subjects)} subjects to delete")
        for i, subject in enumerate(subjects, 1):
            logger.info(f"\nDeleting subject {i} of {len(subjects)}: {subject[1]}")
            if delete_subject(subject[0]):
                logger.info(f"✓ Successfully deleted subject {i}")
            else:
                logger.warning(f"⚠️ Failed to delete subject {i}")
        
        # Delete from database
        execute_db_query(
            "DELETE FROM courses WHERE id = %s",
            (course_id,)
        )
        logger.info("✓ Course record deleted from database")
        logger.info("=== Course deletion completed ===\n")
        return True
    except Exception as e:
        logger.error(f"❌ Error deleting course {course_id}: {str(e)}")
        return False

def ensure_storage_bucket_exists(bucket_name="templates", public=True):
    """
    Ensure that the storage bucket exists and has the correct permissions
    
    Args:
        bucket_name: Name of the bucket to create
        public: Whether the bucket should be public
    
    Returns:
        bool: True if the bucket exists or was created successfully
    """
    try:
        # Use admin client if available for bucket operations
        storage_client = supabase_admin if supabase_admin else supabase
        
        # Check if bucket exists
        try:
            buckets = storage_client.storage.list_buckets()
            bucket_exists = any(bucket.name == bucket_name for bucket in buckets)
            
            if not bucket_exists:
                logger.info(f"Creating new bucket: {bucket_name}")
                try:
                    # Create the bucket
                    storage_client.storage.create_bucket(bucket_name, {'public': public})
                    logger.info(f"Successfully created bucket: {bucket_name}")
                    
                    # If we have admin access, try to set up RLS policies
                    if supabase_admin:
                        try:
                            # Create RLS policies for the bucket
                            logger.info("Setting up RLS policies for the bucket")
                            
                            # Allow all operations for authenticated users
                            # Note: This is a simplified approach. In production, you'd want more granular policies.
                            sql = f"""
                            BEGIN;
                            -- Policy for SELECT (read) operations
                            CREATE POLICY "Allow public access" ON storage.objects
                                FOR SELECT USING (bucket_id = '{bucket_name}');
                                
                            -- Policy for INSERT (upload) operations
                            CREATE POLICY "Allow uploads" ON storage.objects
                                FOR INSERT WITH CHECK (bucket_id = '{bucket_name}');
                                
                            -- Policy for UPDATE operations
                            CREATE POLICY "Allow updates" ON storage.objects
                                FOR UPDATE USING (bucket_id = '{bucket_name}');
                                
                            -- Policy for DELETE operations
                            CREATE POLICY "Allow deletions" ON storage.objects
                                FOR DELETE USING (bucket_id = '{bucket_name}');
                            COMMIT;
                            """
                            
                            # Execute the SQL directly using the Postgres connection
                            # This is a workaround since Supabase Python client doesn't have direct RLS policy management
                            cursor.execute(sql)
                            conn.commit()
                            logger.info("Successfully set up RLS policies for the bucket")
                        except Exception as e:
                            logger.warning(f"Failed to set up RLS policies: {str(e)}")
                except Exception as e:
                    logger.warning(f"Failed to create bucket: {str(e)}. Will attempt to use it anyway if it exists.")
            else:
                logger.info(f"Bucket '{bucket_name}' already exists")
            
            return True
        except Exception as e:
            logger.warning(f"Error checking buckets: {str(e)}. Will attempt to use '{bucket_name}' bucket anyway.")
            return False
    except Exception as e:
        logger.error(f"Failed to ensure bucket exists: {str(e)}")
        return False

def main():
    st.set_page_config(page_title="RAG Course Assessment System", layout="wide")

    # Initialize session state for connections
    if 'pinecone_index' not in st.session_state:
        st.session_state.pinecone_index = init_pinecone()
    
    if 'postgres_conn' not in st.session_state:
        new_conn, new_cursor = init_postgres()
        st.session_state.postgres_conn = new_conn
        st.session_state.postgres_cursor = new_cursor
    
    if 'supabase_client' not in st.session_state or 'supabase_admin' not in st.session_state:
        st.session_state.supabase_client, st.session_state.supabase_admin = init_supabase()

    # Use the persistent connections
    global conn, cursor, index, supabase, supabase_admin
    conn = st.session_state.postgres_conn
    cursor = st.session_state.postgres_cursor
    index = st.session_state.pinecone_index
    supabase = st.session_state.supabase_client
    supabase_admin = st.session_state.supabase_admin
    
    # Ensure the templates bucket exists
    if 'bucket_initialized' not in st.session_state:
        try:
            bucket_exists = ensure_storage_bucket_exists("templates", public=True)
            st.session_state.bucket_initialized = bucket_exists
            if bucket_exists:
                logger.info("Templates bucket is ready for use")
            else:
                logger.warning("Templates bucket may not be properly set up")
        except Exception as e:
            logger.error(f"Error initializing templates bucket: {str(e)}")
            st.session_state.bucket_initialized = False

    # Initialize other session state variables
    if "current_course" not in st.session_state:
        st.session_state.current_course = None
    if "current_subject" not in st.session_state:
        st.session_state.current_subject = None

    # Sidebar navigation
    with st.sidebar:
        st.title("Navigation")
        page = st.radio(
            "Select Page",
            [
                "Course Management",
                "Subject Management",
                "Question Management",
                "Answer Generation",
            ],
        )

    if page == "Course Management":
        show_course_management()
    elif page == "Subject Management":
        show_subject_management()
    elif page == "Question Management":
        show_question_management()
    else:
        show_answer_generation()


def show_course_management():
    st.title("Course Management")

    # Create new course
    with st.form(key="course_form"):
        st.subheader("Create New Course")
        course_name = st.text_input("Course Name")
        submit_course = st.form_submit_button("Create Course")

        if submit_course and course_name:
            course_id = str(uuid.uuid4())
            cursor.execute(
                "INSERT INTO courses (id, name, user_id) VALUES (%s, %s, %s)",
                (course_id, course_name, user_id),
            )
            conn.commit()
            st.success(f"Course '{course_name}' created successfully!")

    # Display existing courses
    st.subheader("Existing Courses")
    cursor.execute("SELECT * FROM courses WHERE user_id = %s", (user_id,))
    courses = cursor.fetchall()
    if courses:
        for course in courses:
            col1, col2, col3 = st.columns([2, 1, 1])
            with col1:
                st.write(f"📚 {course[1]}")
            with col2:
                if st.button("Select", key=f"select_course_{course[0]}"):
                    st.session_state.current_course = course[0]
                    st.success(f"Selected course: {course[1]}")
            with col3:
                if st.button("Delete", key=f"delete_course_{course[0]}"):
                    if delete_course(course[0]):
                        st.success(f"Course '{course[1]}' deleted successfully!")
                        st.session_state.current_course = None
                        st.rerun()
                    else:
                        st.error(f"Failed to delete course '{course[1]}'")
    else:
        st.info("No courses created yet.")


def show_subject_management():
    st.title("Subject Management")

    if not st.session_state.current_course:
        st.warning("Please select a course first from Course Management.")
        return

    # Create new subject
    with st.form(key="subject_form"):
        st.subheader("Add New Subject")
        subject_name = st.text_input("Subject Name")
        guideline_files = st.file_uploader(
            "Upload Guideline Files",
            accept_multiple_files=True,
            type=["pdf", "docx", "txt"],
        )
        submit_subject = st.form_submit_button("Add Subject")

        if submit_subject and subject_name:
            subject_id = str(uuid.uuid4())
            logger.info(f"Creating new subject: {subject_name} (ID: {subject_id})")
            
            # Process guideline files
            if guideline_files:
                documents = []
                for file in guideline_files:
                    try:
                        content = process_uploaded_file(file)
                        documents.append(content)
                        logger.info(f"Successfully processed guideline file: {file.name}")
                    except Exception as e:
                        logger.error(f"Failed to process file {file.name}: {str(e)}")
                        st.error(f"Error processing file {file.name}: {str(e)}")
                        continue
                
                if documents:
                    try:
                        # Create vector store with unique namespace for this subject
                        namespace = f"subject_{subject_id}"
                        vector_store = create_vector_store(documents, namespace)
                        vector_db_url = namespace

                        # Insert into database
                        execute_db_query(
                            "INSERT INTO subjects (id, name, course_id, vector_db_url) VALUES (%s, %s, %s, %s)",
                            (subject_id, subject_name, st.session_state.current_course, vector_db_url)
                        )
                        logger.info(f"Successfully created subject {subject_name} with ID {subject_id}")
                        st.success(f"Subject '{subject_name}' added successfully!")
                    except Exception as e:
                        logger.error(f"Failed to create subject: {str(e)}")
                        st.error(f"Failed to create subject: {str(e)}")
                else:
                    logger.warning("No files were successfully processed")
                    st.error("No files were successfully processed. Please try again.")

    # Display existing subjects
    st.subheader("Existing Subjects")
    subjects = execute_db_query(
        "SELECT * FROM subjects WHERE course_id = %s",
        (st.session_state.current_course,),
        fetch=True
    )
    logger.info(f"Retrieved {len(subjects)} subjects for course {st.session_state.current_course}")
    if subjects:
        for subject in subjects:
            col1, col2, col3 = st.columns([2, 1, 1])
            with col1:
                st.write(f"📘 {subject[1]}")
            with col2:
                if st.button("Select", key=f"select_subject_{subject[0]}"):
                    st.session_state.current_subject = subject[0]
                    st.success(f"Selected subject: {subject[1]}")
            with col3:
                if st.button("Delete", key=f"delete_subject_{subject[0]}"):
                    if delete_subject(subject[0]):
                        st.success(f"Subject '{subject[1]}' deleted successfully!")
                        st.session_state.current_subject = None
                        st.rerun()
                    else:
                        st.error(f"Failed to delete subject '{subject[1]}'")
    else:
        st.info("No subjects created yet.")


def show_question_management():
    st.title("Question Management")

    if not st.session_state.current_subject:
        st.warning("Please select a subject first from Subject Management.")
        return

    # Initialize the guideline input method in session state if not present
    if "guideline_input_method" not in st.session_state:
        st.session_state.guideline_input_method = "Upload Files"
    if "sample_answer_input_method" not in st.session_state:
        st.session_state.sample_answer_input_method = "Upload Files"

    col1, col2 = st.columns(2)
    with col1:
    # Add radio button outside the form for instant updates
        selected_method_for_guideline = st.radio(
            "How would you like to provide guidelines?",
            ["Upload Files", "Text Input"],
            key="guideline",
        )
    with col2:
        selected_method_for_sample_answers = st.radio(
            "How would you like to provide sample answers?",
            ["Upload Files", "Text Input"],
            key="sample_answer",
        )
    # Update session state and log the selection
    st.session_state.guideline_input_method = selected_method_for_guideline
    logger.info(f"Selected guideline input method: {selected_method_for_guideline}")

    st.session_state.sample_answer_input_method = selected_method_for_sample_answers
    logger.info(
        f"Selected guideline input method: {selected_method_for_sample_answers}"
    )

    # Create new question
    with st.form(key="question_form"):
        st.subheader("Add New Question")
        question_text = st.text_area("Question Text")

        # Initialize variables
        guideline_files = None
        guideline_text = None
        sample_answers_files = None
        sample_answers_text = None

        # Show only the selected input method based on session state
        if st.session_state.guideline_input_method == "Upload Files":
            logger.debug("Showing file upload fields")
            guideline_files = st.file_uploader(
                "Upload Guideline Files",
                accept_multiple_files=True,
                type=["pdf", "docx", "txt"],
            )
        else:  # Text Input
            logger.debug("Showing text input fields")
            guideline_text = st.text_area(
                "Enter Guidelines",
                help="Enter the guidelines for this question directly as text.",
            )

        if st.session_state.sample_answer_input_method == "Upload Files":
            sample_answers_files = st.file_uploader(
                "Upload Sample Answers",
                accept_multiple_files=True,
                type=["pdf", "docx", "txt"],
            )
        else:
            sample_answers_text = st.text_area(
                "Enter Sample Answers",
                help="Enter the sample answers for this question directly as text.",
            )

        # Add template file upload
        template_file = st.file_uploader(
            "Upload Answer Template (Optional)",
            accept_multiple_files=False,
            type=["docx"],
            help="Upload a DOCX template that will be used when generating answers. The template structure will be preserved in the answer."
        )

        instructions = st.text_area("Additional Instructions")
        submit_question = st.form_submit_button("Add Question")

        if submit_question and question_text:
            with st.spinner("Please wait, it may take some time"):
                question_id = str(uuid.uuid4())
                
                # Process guidelines
                guideline_vector_db_url, guideline_text_content = handle_content_input(
                    input_method=st.session_state.guideline_input_method,
                    files=guideline_files,
                    text_content=guideline_text,
                    question_id=question_id,
                    content_type="guidelines"
                )
            
                # Process sample answers
                sample_answers_vector_db_url, sample_answers_text_content = handle_content_input(
                    input_method=st.session_state.sample_answer_input_method,
                    files=sample_answers_files,
                    text_content=sample_answers_text,
                    question_id=question_id,
                    content_type="sample_answers"
                )
                
                # Process template file
                template_url = None
                if template_file:
                    try:
                        template_url = upload_template_to_supabase(template_file, question_id)
                        st.success(f"Template file uploaded successfully!")
                    except Exception as e:
                        logger.error(f"Failed to upload template file: {str(e)}")
                        st.error(f"Failed to upload template file: {str(e)}")
            
                # Set boolean flags based on input methods
                is_guideline_textbox = st.session_state.guideline_input_method == "Text Input"
                is_sample_answer_textbox = st.session_state.sample_answer_input_method == "Text Input"
                
                # Before database insert
                logger.info("=== Database Operation ===")
                logger.info(f"Guideline vector DB URL: {guideline_vector_db_url}")
                logger.info(f"Sample answers vector DB URL: {sample_answers_vector_db_url}")
                logger.info(f"Template URL: {template_url}")
                cursor.execute(
                    """INSERT INTO questions 
                    (id, text, guideline_vector_db_url, guideline_text_box, is_guideline_textbox, 
                        sample_answers_vector_db_url, sample_answers_textbox, is_sample_answer_textbox, 
                        instructions, template_url, subject_id) 
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING is_guideline_textbox""",
                    (
                        question_id,
                        question_text,
                        guideline_vector_db_url,  # Will be None for text input
                        guideline_text_content,  # Will be None for file upload
                        is_guideline_textbox,  # True for text input, False for file upload
                        sample_answers_vector_db_url,  # Will be None for text input
                        sample_answers_text_content, # Will be None for file upload
                        is_sample_answer_textbox,
                        instructions,
                        template_url,
                        st.session_state.current_subject,
                    ),
                )
                stored_value = cursor.fetchone()[0]
                logger.info(f"Value stored in database: {stored_value} ({type(stored_value)})")
                conn.commit()
                st.success("Question added successfully!")

    # Display existing questions
    st.subheader("Existing Questions")
    cursor.execute(
        """
        SELECT 
            q.*, 
            to_char(q.created_at, 'YYYY-MM-DD') as formatted_date 
        FROM questions q 
        WHERE subject_id = %s
        """,
        (st.session_state.current_subject,),
    )
    questions = cursor.fetchall()

    if questions:
        for question in questions:
            question_text = question[1]
            instructions = question[8]
            guideline_text = question[3]
            sample_answers_text = question[6]
            sample_answers_file = question[5]
            is_guideline_textbox = bool(question[4])
            is_sample_answer_textbox = bool(question[7])
            template_url = question[9]  # New template_url column
            created_date = question[-1]

            with st.expander(f"📝 {question_text}"):
                # Display instructions if present
                if instructions:
                    st.write("Instructions:", instructions)
                
                # Display guidelines
                guideline_label = "Guidelines:"
                if is_guideline_textbox:
                    st.write(f"{guideline_label} {guideline_text}")
                else:
                    st.write(f"{guideline_label} [File Upload]")
                
                # Display sample answers if present
                sample_label = "Sample Answers:"
                if is_sample_answer_textbox and sample_answers_text:
                    st.write(f"{sample_label} {sample_answers_text}")
                elif not is_sample_answer_textbox and sample_answers_file:
                    st.write(f"{sample_label} [File Upload]")
                
                # Display template information
                if template_url:
                    st.write(f"Template: [Available]({template_url})")
                
                # Display creation date
                st.write(f"Created At: {created_date}")

                # Add delete button
                if st.button("Delete Question", key=f"delete_question_{question[0]}"):
                    if delete_question(question[0]):
                        st.success("Question deleted successfully!")
                        st.rerun()
                    else:
                        st.error("Failed to delete question")
    else:
        st.info("No questions added yet.")


def show_answer_generation():
    st.title("Answer Generation")

    if not st.session_state.current_subject:
        st.warning("Please select a subject first from Subject Management.")
        return

    # Student information
    col1, col2 = st.columns(2)
    with col1:
        student_name = st.text_input("Student Name")
    with col2:
        student_id = st.text_input("Student ID")

    # Generate all answers button
    if st.button("Generate All Answers"):
        if student_name and student_id:
            try:
                with st.spinner("Generating answers for all questions... This may take a while."):
                    # Initialize OpenAI embeddings and LLM
                    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
                    llm = ChatOpenAI(model="gpt-4", temperature=0.2)
                    
                    # Generate answers for all questions
                    all_qa_pairs = generate_all_answers(
                        st.session_state.current_subject,
                        embeddings,
                        llm
                    )
                    
                    if all_qa_pairs:
                        st.success("Successfully generated all answers!")
                        
                        # Create temporary file for all Q&As
                        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_all_qa:
                            create_all_qa_doc(all_qa_pairs, tmp_all_qa.name)
                            with open(tmp_all_qa.name, 'rb') as docx_file:
                                docx_bytes = docx_file.read()
                            st.download_button(
                                "Download All Q&As (DOCX)",
                                data=docx_bytes,
                                file_name=f"all_qa_{student_id}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            
                            # Cleanup temporary file
                            try:
                                os.unlink(tmp_all_qa.name)
                            except Exception as e:
                                logger.warning(f"Error cleaning up temporary file: {str(e)}")
                    else:
                        st.warning("No questions found for the current subject.")
            except Exception as e:
                error_msg = f"Error generating answers: {str(e)}"
                logger.error(error_msg)
                st.error(error_msg)
        else:
            st.warning("Please enter student name and ID.")


if __name__ == "__main__":
    main()
